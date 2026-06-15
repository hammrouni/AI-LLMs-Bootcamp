"""
Generates the sample files used by demo.py PART 3:
  sample_docs/handbook.pdf   (3 pages, with headers/page numbers/hyphenation)
  sample_docs/policies.docx  (messy spacing, repeated header paragraph)
  sample_docs/notes.txt      (plain text)

Run once:
    pip install reportlab python-docx
    python generate_sample_docs.py
"""

import os

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

OUT_DIR = os.path.join(os.path.dirname(__file__), "sample_docs")
os.makedirs(OUT_DIR, exist_ok=True)


def make_pdf():
    path = os.path.join(OUT_DIR, "handbook.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4

    pages = [
        [
            "MASTER Soft — Internal — Confidential",
            "Page 1 / 3",
            "",
            "Welcome to the MASTER Soft employee hand-",
            "book. This document explains our company poli-",
            "cies for our Tunis and Sousse offices.",
        ],
        [
            "MASTER Soft — Internal — Confidential",
            "Page 2 / 3",
            "",
            "Refund policy: Customers can request a re-",
            "fund within 14 days of purchase, no questions",
            "asked. Contact Yasmine in customer service.",
        ],
        [
            "MASTER Soft — Internal — Confidential",
            "Page 3 / 3",
            "",
            "Vacation:    Employees    accrue",
            "2 days per month, totalling 24 days/year.",
        ],
    ]

    for page_lines in pages:
        y = height - 72
        for line in page_lines:
            c.drawString(72, y, line)
            y -= 18
        c.showPage()

    c.save()
    print(f"Wrote {path}")


def make_docx():
    path = os.path.join(OUT_DIR, "policies.docx")
    doc = Document()
    doc.add_paragraph("MASTER Soft — Internal — Confidential")
    doc.add_paragraph(
        "Refund   policy:   Customers  can  request  a  refund  within  14  days  "
        "of  purchase,  no  questions  asked.   Contact  Yasmine  in  customer  service."
    )
    doc.add_paragraph(
        "Vacation:    Employees    accrue   2 days per month, totalling 24 days/year."
    )
    doc.save(path)
    print(f"Wrote {path}")


def make_txt():
    path = os.path.join(OUT_DIR, "notes.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(
            "MASTER Soft — Internal — Confidential\n"
            "Page 1 / 1\n\n"
            "Working hours are 9:00 to 18:00, Monday to Fri-\n"
            "day. Remote work is allowed up to 2 days per\n"
            "week with manager approval.\n"
        )
    print(f"Wrote {path}")


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_txt()
